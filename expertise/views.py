# expertise/views.py

from django.shortcuts import get_object_or_404, render, redirect
from django.urls import reverse_lazy, reverse
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.db.models import Q, Sum
from django.db import transaction
from django.http import HttpResponse
from django.utils import timezone
from datetime import datetime
from num2words import num2words
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import barcode
import io
import base64
from barcode.writer import ImageWriter
from .models import FicheEvenement, PersonnelNavigant, CompagnieAerienne, Bordereau, FactureMedecin
from .forms import BordereauSelectionForm
from django.db import models
from django.template.loader import render_to_string
from weasyprint import HTML



def assign_bordereau(request, mois, annee, iata):
    compagnie = get_object_or_404(CompagnieAerienne, iata=iata)
    evenements = FicheEvenement.objects.filter(
        date_evenement__year=annee, date_evenement__month=mois, personnel__compagnie=compagnie
    )

    # Création du bordereau
    date_creation = datetime.today()
    no_bordereau = f"EB{date_creation.day:02d}{mois:02d}{str(annee)[-2:]}{iata}"
    bordereau = Bordereau.objects.create(
        date_bordereau=date_creation,
        no_bordereau=no_bordereau
    )

    # Lier le bordereau aux événements
    for evenement in evenements:
        evenement.no_bordereau = no_bordereau
        evenement.bordereau = bordereau
        evenement.save()

    return redirect('bordereau_view', mois=mois, annee=annee, iata=iata)

# ----- VUES POUR LES PERSONNELS -----
class PersonnelListView(ListView):
    model = PersonnelNavigant
    template_name = 'expertise/personnel_list.html'
    context_object_name = 'personnels'

    def get_queryset(self):
        queryset = super().get_queryset()
        query = self.request.GET.get('q')
        if query:
            queryset = queryset.filter(
                Q(nom__icontains=query) |
                Q(prenom__icontains=query) |
                Q(dn__icontains=query)
            )
        return queryset


class PersonnelDetailView(DetailView):
    model = PersonnelNavigant
    template_name = 'expertise/personnel_detail.html'
    context_object_name = 'personnel'
    slug_field = 'dn'
    slug_url_kwarg = 'dn'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['evenements'] = self.object.evenements.all()
        return context


# ----- VUES POUR LES EVENEMENTS -----
from django.utils.timezone import now

class FicheEvenementCreateView(CreateView):
    model = FicheEvenement
    fields = [
        'date_evenement',
        'cs_cempn', 'date_cempn', 'honoraire_cempn', 'medecin_cempn',
        'cs_oph', 'date_cs_oph', 'honoraire_cs_oph', 'medecin_oph',
        'cs_orl', 'date_cs_orl', 'honoraire_cs_orl', 'medecin_orl',
        'cs_labo', 'date_cs_labo', 'honoraire_cs_labo', 'medecin_labo',
        'cs_lbx', 'date_cs_lbx', 'honoraire_cs_lbx',
        'cs_toxique', 'date_cs_toxique', 'honoraire_cs_toxique',
        'cs_radio', 'date_cs_radio', 'honoraire_cs_radio', 'medecin_radio',
        'frais_dossier', 'quote_part_patient',
        'paiement', 'date_paiement', 'modalite_paiement',
    ]
    template_name = 'expertise/evenement_form.html'

    def get_initial(self):
        date = now().date()  # ou une autre logique si tu veux personnaliser
        return {
            'date_evenement': date,
            'date_cempn': date,
            'date_cs_oph': date,
            'date_cs_orl': date,
            'date_cs_labo': date,
            'date_cs_lbx': date,
            'date_cs_toxique': date,
            'date_cs_radio': date,
            'date_paiement': date
        }

    def form_valid(self, form):
        dn = self.kwargs['dn']
        personnel = get_object_or_404(PersonnelNavigant, dn=dn)
        form.instance.personnel = personnel
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['dn'] = self.kwargs['dn']
        return context

    def get_success_url(self):
        return reverse_lazy('personnel_detail', kwargs={'dn': self.object.personnel.dn})




class FactureView(DetailView):
    model = FicheEvenement
    template_name = 'expertise/facture.html'
    context_object_name = 'evenement'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        invoice_number = self.object.no_facture
        print(f"DEBUG: Numéro de facture = {invoice_number}")

        if invoice_number:
            try:
                Code128 = barcode.get_barcode_class('code128')
                barcode_instance = Code128(invoice_number, writer=ImageWriter())
                buffer = io.BytesIO()
                barcode_instance.write(buffer)
                barcode_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                context['barcode'] = barcode_base64
            except Exception as e:
                print(f"Erreur lors de la génération du code-barres : {e}")

        return context


# ----- VUE DU BORDEREAU -----
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.db import transaction
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from num2words import num2words

from .models import CompagnieAerienne, FicheEvenement, Bordereau

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.db import transaction
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words

from .models import CompagnieAerienne, FicheEvenement, Bordereau
from docx.shared import RGBColor

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.db import transaction
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from num2words import num2words

from .models import CompagnieAerienne, FicheEvenement, Bordereau


def download_bordereau(request, mois, annee, iata):
    compagnie = get_object_or_404(CompagnieAerienne, iata=iata)
    evenements = FicheEvenement.objects.filter(
        date_evenement__year=annee,
        date_evenement__month=mois,
        personnel__compagnie=compagnie
    )

    date_creation = datetime.today()
    no_bordereau = f"EB{date_creation.day:02d}{mois:02d}{str(annee)[-2:]}{iata}"

    bordereau, created = Bordereau.objects.get_or_create(
        no_bordereau=no_bordereau,
        defaults={"date_bordereau": date_creation}
    )

    with transaction.atomic():
        evenements.update(bordereau=bordereau)

    doc = Document()
    para = doc.add_heading('Centre Médical du Personnel Naviguant de Polynésie française', level=1)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('Dr. Christian Hellec', level=1)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('BP 1946 - Papeete - Tahiti', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('Polynésie Française', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('Bordereau de dépôt de factures', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('--------------------', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Date de création : {date_creation.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Numéro du bordereau : {no_bordereau}")
    doc.add_paragraph(f"Compagnie aérienne : {compagnie.nom} ({compagnie.iata})")

    # Tableau principal
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdrs = ["Numéro de facture", "DN", "Nom", "Prénom", "Total (XPF)", "Paiement"]
    for i, header in enumerate(hdrs):
        cell = table.rows[0].cells[i]
        cell.text = header
        para = cell.paragraphs[0]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.runs[0]
        run.bold = True

    total_general = 0

    for e in evenements:
        total_general += e.total or 0
        row = table.add_row().cells
        data = [
            e.no_facture or "N/A",
            e.personnel.dn,
            e.personnel.nom,
            e.personnel.prenom,
            f"{e.total or 0:,} XPF",
            "Payé" if e.paiement else "Non payé"
        ]
        for i, val in enumerate(data):
            para = row[i].paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.add_run(val)

    total_lettres = num2words(total_general, lang='fr').capitalize()
    doc.add_paragraph(
        f"\nNombre de factures : {evenements.count()} | Total général : {total_general:,} XPF ({total_lettres})"
    )
    doc.add_paragraph("Dr. Christian HELLEC")

    # Factures individuelles
    for e in evenements:
        doc.add_page_break()
        para = doc.add_heading('Centre Médical du Personnel Naviguant de POlynésie française', level=1)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading('Dr. Christian Hellec', level=1)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading('BP 1946 - Papeete - Tahiti', level=2)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading('Polynésie Française', level=2)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading('Facture Individuelle', level=2)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading('--------------------', level=2)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("")
        doc.add_paragraph(f"Date : {e.date_evenement.strftime('%d/%m/%Y')}")
        doc.add_paragraph(f"Numéro de facture : {e.no_facture or 'N/A'}")

        doc.add_heading("Informations du patient", level=2)
        doc.add_paragraph(f"Nom : {e.personnel.nom}")
        doc.add_paragraph(f"Prénom : {e.personnel.prenom}")
        doc.add_paragraph(f"DN : {e.personnel.dn}")
        if e.personnel.date_de_naissance:
            doc.add_paragraph(f"Date de naissance : {e.personnel.date_de_naissance.strftime('%d/%m/%Y')}")

        doc.add_heading("Détails des actes", level=2)
        actes = [
            ("CEMPN/Pf", e.cs_cempn, e.date_cempn, e.medecin_cempn, e.honoraire_cempn),
            ("Ophtalmologie", e.cs_oph, e.date_cs_oph, e.medecin_oph, e.honoraire_cs_oph),
            ("ORL", e.cs_orl, e.date_cs_orl, e.medecin_orl, e.honoraire_cs_orl),
            ("Biologie sanguine", e.cs_labo, e.date_cs_labo, e.medecin_labo, e.honoraire_cs_labo),
            ("Biologie urinaire", e.cs_lbx, e.date_cs_lbx, e.medecin_labo, e.honoraire_cs_lbx),
            ("Toxicologie", e.cs_toxique, e.date_evenement, e.medecin_labo, e.honoraire_cs_toxique),
            ("Radiologie", e.cs_radio, e.date_cs_radio, e.medecin_radio, e.honoraire_cs_radio),
        ]

        # Optionnel si ce champ existe dans ton modèle :
        #if hasattr(e, 'honoraire_cs_toxique'):
            #actes.append(("TOXICO", True, e.date_evenement, e.medecin_labo, e.honoraire_cs_toxique))

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        headers = ["Date", "Acte", "Médecin", "Montant (XPF)", "Quote-part"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            para = cell.paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.runs[0].bold = True

        for libelle, actif, date_acte, medecin, montant in actes:
            if actif:
                row = table.add_row().cells
                data = [
                    date_acte.strftime('%d/%m/%Y') if date_acte else "-",
                    libelle,
                    f"{medecin.prenom} {medecin.nom}" if medecin else "-",
                    f"{montant or 0:,} XPF",
                    f"{e.paye_par_patient or 0:,} XPF" if e.quote_part_patient else "-"
                ]
                for i, val in enumerate(data):
                    para = row[i].paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    para.add_run(val)

        doc.add_paragraph(f"\n💰 Total : {e.total or 0:,} XPF")
        doc.add_paragraph(f"🧾 Payé par le patient : {e.paye_par_patient or 0:,} XPF")
        doc.add_paragraph("Dr. Christian HELLEC")

    # Générer le fichier en réponse HTTP
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Bordereau_{no_bordereau}.docx"'
    doc.save(response)
    return response



# ----- LISTE DES BORDEREAUX -----
from django.shortcuts import render
from .models import Bordereau
from django.db.models import Sum
from .models import Bordereau

def liste_bordereaux(request):
    bordereaux = Bordereau.objects.all().prefetch_related('evenements')
    
    for bordereau in bordereaux:
        bordereau.total_general = (
            bordereau.evenements.aggregate(Sum('total'))['total__sum'] or 0
        )
        bordereau.nb_factures = bordereau.evenements.count()

    return render(request, 'expertise/liste_bordereaux.html', {
        'bordereaux': bordereaux
    })





# ----- SELECTION DE BORDEREAU -----
def bordereau_selection_view(request):
    if request.method == 'POST':
        form = BordereauSelectionForm(request.POST)
        if form.is_valid():
            mois = form.cleaned_data['mois']
            annee = form.cleaned_data['annee']
            compagnie = form.cleaned_data['compagnie']
            iata = compagnie.iata
            return redirect('bordereau_detail', annee=annee, mois=mois, iata=iata)
    else:
        form = BordereauSelectionForm()

    return render(request, 'expertise/selection_bordereau.html', {'form': form})

# expertise/views.py

from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse_lazy, reverse
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.db.models import Q
from .models import PersonnelNavigant, FicheEvenement, CompagnieAerienne
from .forms import BordereauSelectionForm
from .utils import nombre_en_lettres
from num2words import num2words
from datetime import datetime
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from django.db import transaction
import io
import base64
import barcode
from barcode.writer import ImageWriter

# ----- VUES POUR LES PERSONNELS -----

class PersonnelListView(ListView):
    model = PersonnelNavigant
    template_name = 'expertise/personnel_list.html'
    context_object_name = 'personnels'

    def get_queryset(self):
        queryset = super().get_queryset()
        query = self.request.GET.get('q')
        if query:
            queryset = queryset.filter(
                Q(nom__icontains=query) |
                Q(prenom__icontains=query) |
                Q(dn__icontains=query)
            )
        return queryset


class PersonnelDetailView(DetailView):
    model = PersonnelNavigant
    template_name = 'expertise/personnel_detail.html'
    context_object_name = 'personnel'
    slug_field = 'dn'
    slug_url_kwarg = 'dn'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['evenements'] = self.object.evenements.all()
        return context


class PersonnelCreateView(CreateView):
    model = PersonnelNavigant
    fields = ['dn', 'nom', 'prenom', 'date_de_naissance', 'sexe', 'statut_pn', 'compagnie']
    template_name = 'expertise/personnel_form.html'
    success_url = reverse_lazy('personnel_list')


class PersonnelUpdateView(UpdateView):
    model = PersonnelNavigant
    fields = ['dn', 'nom', 'prenom', 'date_de_naissance', 'sexe', 'statut_pn', 'compagnie']
    template_name = 'expertise/personnel_form.html'
    slug_field = 'dn'
    slug_url_kwarg = 'dn'
    success_url = reverse_lazy('personnel_list')


class PersonnelDeleteView(DeleteView):
    model = PersonnelNavigant
    template_name = 'expertise/personnel_confirm_delete.html'
    slug_field = 'dn'
    slug_url_kwarg = 'dn'
    success_url = reverse_lazy('personnel_list')

from django.views.generic import UpdateView
from .models import FicheEvenement
from django.urls import reverse_lazy

class FicheEvenementUpdateView(UpdateView):
    model = FicheEvenement
    fields = [
        'date_evenement',
        'cs_cempn', 'date_cempn', 'honoraire_cempn',
        'cs_oph', 'date_cs_oph', 'honoraire_cs_oph',
        'cs_orl', 'date_cs_orl', 'honoraire_cs_orl',
        'cs_labo', 'date_cs_labo', 'honoraire_cs_labo',
        'cs_lbx', 'date_cs_lbx', 'honoraire_cs_lbx',
        'cs_toxique', 'date_cs_toxique', 'honoraire_cs_toxique',
        'cs_radio', 'date_cs_radio', 'honoraire_cs_radio',
        'medecin_cempn', 'medecin_oph',
        'medecin_orl', 'medecin_radio', 'medecin_labo',
        'frais_dossier', 'quote_part_patient',
        'paiement', 'date_paiement', 'modalite_paiement',
    ]
    template_name = 'expertise/evenement_form.html'

    def get_success_url(self):
        if self.object.personnel and self.object.personnel.dn:
            return reverse_lazy('personnel_detail', kwargs={'dn': self.object.personnel.dn})
        else:
            return reverse_lazy('personnel_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.object.personnel:
            context['dn'] = self.object.personnel.dn
        return context


from django.views.generic import DeleteView
from django.urls import reverse_lazy
from .models import FicheEvenement

class FicheEvenementDeleteView(DeleteView):
    model = FicheEvenement
    template_name = 'expertise/evenement_confirm_delete.html'

def get_success_url(self):
    personnel = getattr(self.object, 'personnel', None)
    if personnel:
        return reverse_lazy('personnel_detail', kwargs={'dn': personnel.dn})
    else:
        return reverse_lazy('personnel_list')  # Fallback si pas de personnel


from django.shortcuts import render, get_object_or_404
from datetime import datetime
from .models import FicheEvenement, CompagnieAerienne
from .utils import nombre_en_lettres  # Assure-toi que ce module existe

def bordereau_view(request, annee, mois, iata):
    compagnie = get_object_or_404(CompagnieAerienne, iata=iata)

    evenements = FicheEvenement.objects.filter(
        date_evenement__year=annee,
        date_evenement__month=mois,
        personnel__compagnie=compagnie
    )

    date_bordereau = datetime.today().strftime('%d/%m/%Y')
    no_bordereau = f"EB{datetime.today().day:02d}{mois:02d}{str(annee)[-2:]}{iata}"
    total_global = sum(e.total for e in evenements)
    total_global_lettres = nombre_en_lettres(total_global)

    # Création ou récupération du bordereau
    bordereau, _ = Bordereau.objects.get_or_create(
    no_bordereau=no_bordereau,
    defaults={'date_bordereau': datetime.today()}
)

    # Mise à jour de chaque événement
    with transaction.atomic():
        for e in evenements:
            e.bordereau = bordereau
            e.save()


    return render(request, "expertise/bordereau.html", {
        "evenements": evenements,
        "mois": mois,
        "annee": annee,
        "iata": iata,
        "compagnie": compagnie,
        "date_bordereau": date_bordereau,
        "no_bordereau": no_bordereau,
        "nombre_factures": evenements.count(),
        "total_global": total_global,
        "total_global_lettres": total_global_lettres,
    })
# views.py

from .models import FicheEvenement

def bordereau_factures(request, no_bordereau):
    evenements = FicheEvenement.objects.filter(bordereau__no_bordereau=no_bordereau)
    return render(request, 'expertise/factures_bordereau.html', {
        'evenements': evenements,
        'no_bordereau': no_bordereau
    })

from django.shortcuts import redirect, get_object_or_404
from .models import Bordereau
from django.views.decorators.csrf import csrf_exempt
from django.http import HttpResponseRedirect

@csrf_exempt
#from django.db.models import Sum
#from .models import FactureMedecin  # n'oublie pas d'importer

def toggle_virement(request, id):
    bordereau = get_object_or_404(Bordereau, id=id)
    bordereau.virement = not bordereau.virement
    bordereau.save()

    # 💡 Création des factures pour les médecins si virement = True
    if bordereau.virement:
        evenements = bordereau.evenements.select_related(
            'medecin_cempn', 'medecin_oph', 'medecin_orl',
            'medecin_radio', 'medecin_labo'
        )

        honoraires_medecins = {}

        for e in evenements:
            for champ, montant in [
                (e.medecin_cempn, e.honoraire_cempn),
                (e.medecin_oph, e.honoraire_cs_oph),
                (e.medecin_orl, e.honoraire_cs_orl),
                (e.medecin_radio, e.honoraire_cs_radio),
                (e.medecin_labo, e.honoraire_cs_labo),
                (e.medecin_labo, e.honoraire_cs_lbx),
                (e.medecin_labo, e.honoraire_cs_toxique),
            ]:
                if champ:
                    honoraires_medecins.setdefault(champ, 0)
                    honoraires_medecins[champ] += montant or 0

        for medecin, total in honoraires_medecins.items():
            FactureMedecin.objects.create(
                medecin=medecin,
                bordereau=bordereau,
                montant=total
            )

    return redirect('liste_bordereaux')


from decimal import Decimal

def factures_medecins_bordereau(request, no_bordereau):
    bordereau = get_object_or_404(Bordereau, no_bordereau=no_bordereau)
    factures = FactureMedecin.objects.filter(bordereau=bordereau)

    for facture in factures:
        if facture.medecin.nom.lower() == "hellec":
            facture.redevance = Decimal('0')
        else:
            facture.redevance = round(facture.montant * Decimal('0.06'), 0)
        facture.montant_net = round(facture.montant - facture.redevance, 0)

    return render(request, 'expertise/factures_medecins_bordereau.html', {
        'bordereau': bordereau,
        'factures': factures,
    })



from django.shortcuts import get_object_or_404, redirect
from .models import Bordereau

def supprimer_bordereau(request, id):
    bordereau = get_object_or_404(Bordereau, id=id)

    # Dissocier les événements liés
    for evenement in bordereau.evenements.all():
        evenement.bordereau = None
        evenement.save()

    # Supprimer le bordereau
    bordereau.delete()

    return redirect('liste_bordereaux')

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from .models import FactureMedecin, Bordereau, FicheEvenement
from decimal import Decimal

from docx import Document
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from decimal import Decimal
from .models import Bordereau, Medecin, FactureMedecin, FicheEvenement

def telecharger_facture_medecin(request, bordereau_no, medecin_id):
    bordereau = get_object_or_404(Bordereau, no_bordereau=bordereau_no)
    medecin = get_object_or_404(Medecin, id=medecin_id)

    evenements = FicheEvenement.objects.filter(
        bordereau=bordereau
    ).filter(
        models.Q(medecin_cempn=medecin) |
        models.Q(medecin_oph=medecin) |
        models.Q(medecin_orl=medecin) |
        models.Q(medecin_radio=medecin) |
        models.Q(medecin_labo=medecin)
    )

    doc = Document()
    para = doc.add_heading('Centre Médical du Personnel Naviguant de Polynésie française', level=1)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('Dr. Christian Hellec', level=1)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('BP 295 - Papeete - Tahiti', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('Polynésie Française', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('--------------------', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading(f"Facture pour le Dr {medecin.nom} {medecin.prenom}", level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('--------------------', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")
    doc.add_paragraph(f"Bordereau : {bordereau_no}")

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Date"
    hdr_cells[1].text = "Patient"
    hdr_cells[2].text = "Montant brut"
    hdr_cells[3].text = "Redevance"
    hdr_cells[4].text = "Montant net"

    total = Decimal("0.00")
    total_redevance = Decimal("0.00")

    for e in evenements:
        if e.medecin_cempn == medecin:
            montant = e.honoraire_cempn or 0
        elif e.medecin_oph == medecin:
            montant = e.honoraire_cs_oph or 0
        elif e.medecin_orl == medecin:
            montant = e.honoraire_cs_orl or 0
        elif e.medecin_radio == medecin:
            montant = e.honoraire_cs_radio or 0
        elif e.medecin_labo == medecin:
            montant = (e.honoraire_cs_labo or 0) + (e.honoraire_cs_lbx or 0) + (e.honoraire_cs_toxique or 0)
        else:
            continue

        # Patient
        nom_patient = f"{e.personnel.prenom} {e.personnel.nom}"

        # Calculs
        redevance = Decimal("0.00") if medecin.nom.upper() == "HELLEC" else montant * Decimal("0.06")
        net = montant - redevance

        total += montant
        total_redevance += redevance

        row = table.add_row().cells
        row[0].text = str(e.date_evenement)
        row[1].text = nom_patient
        row[2].text = f"{montant:.0f} XPF"
        row[3].text = f"{redevance:.0f} XPF"
        row[4].text = f"{net:.0f} XPF"

    doc.add_paragraph("")
    doc.add_paragraph(f"Total brut : {total:.0f} XPF")
    doc.add_paragraph(f"Total redevance : {total_redevance:.0f} XPF")
    doc.add_paragraph(f"Total net à payer : {total - total_redevance:.0f} XPF")
    if medecin.iban:
        doc.add_paragraph(f"IBAN / RIB : {medecin.iban}")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Facture_{medecin.nom}_{bordereau_no}.docx".replace(" ", "_")
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)

    return response

from decimal import Decimal

def telecharger_facture_medecin(request, bordereau_no, medecin_id):
    bordereau = get_object_or_404(Bordereau, no_bordereau=bordereau_no)
    medecin = get_object_or_404(Medecin, id=medecin_id)

    evenements_originaux = FicheEvenement.objects.filter(
        bordereau=bordereau
    ).filter(
        models.Q(medecin_cempn=medecin) |
        models.Q(medecin_oph=medecin) |
        models.Q(medecin_orl=medecin) |
        models.Q(medecin_radio=medecin) |
        models.Q(medecin_labo=medecin)
    )

    evenements = []
    total_brut = Decimal("0.00")
    total_redevance = Decimal("0.00")

    for e in evenements_originaux:
        if e.medecin_cempn == medecin:
            montant = e.honoraire_cempn or 0
        elif e.medecin_oph == medecin:
            montant = e.honoraire_cs_oph or 0
        elif e.medecin_orl == medecin:
            montant = e.honoraire_cs_orl or 0
        elif e.medecin_radio == medecin:
            montant = e.honoraire_cs_radio or 0
        elif e.medecin_labo == medecin:
            montant = sum(filter(None, [e.honoraire_cs_labo, e.honoraire_cs_lbx, e.honoraire_cs_toxique]))
        else:
            continue

        redevance = Decimal("0.00") if medecin.nom.upper() == "HELLEC" else montant * Decimal("0.06")
        net = montant - redevance

        evenements.append({
            "date": e.date_evenement.strftime('%d/%m/%Y'),
            "patient": f"{e.personnel.prenom} {e.personnel.nom}",
            "montant": f"{montant:.0f}",
            "redevance": f"{redevance:.0f}",
            "net": f"{net:.0f}",
        })

        total_brut += montant
        total_redevance += redevance

    total_net = total_brut - total_redevance

    html_string = render_to_string("expertise/facture_medecin_pdf.html", {
        "medecin": medecin,
        "bordereau": bordereau,
        "evenements": evenements,
        "total_brut": f"{total_brut:.0f}",
        "total_redevance": f"{total_redevance:.0f}",
        "total_net": f"{total_net:.0f}",
    })

    pdf_file = HTML(string=html_string).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Facture_{medecin.nom}_{bordereau.no_bordereau}.pdf"'
    return response
